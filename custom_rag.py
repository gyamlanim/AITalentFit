from llm_interface import CustomLLM
from vector_db import VectorDB
from typing import Dict, List, Any
from PyPDF2 import PdfReader
from docx import Document
import io
import json
import re
from io import BytesIO

class CustomRAG:
    def __init__(self, db_path: str = "resume_db", table_name: str = "resume_table"):
        """Initialize CustomRAG with LLM and VectorDB instances."""
        self.llm = CustomLLM()
        self.vector_db = VectorDB(db_path)
        self.vector_db.create_table(table_name)
        self.db_path = db_path
        self.table_name = table_name

    def read_pdf(self, file_path) -> str:
        text = ""
        if file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text()
        elif file_path.endswith(".docx"):
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        return text.strip()
    
    def read_pdf_bytes(self, file_bytes, file_name) -> str:
        file = BytesIO(file_bytes)
        text = ""
        if file_name.endswith(".pdf"):
            file_extension = ".pdf"
            reader = PdfReader(file)
            for page in reader.pages:
                text += page.extract_text()
        elif file_name.endswith(".docx"):
            file_extension = ".docx"
            doc = Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        return text.strip(), file_extension


    def enter_resume_to_db_save_file(self, file_bytes, file_name) -> str:
        try:
            resume_text, file_extension = self.read_pdf_bytes(file_bytes, file_name)
            summary_response = json.loads(self.summarize_resume(resume_text))
            summary = summary_response.get("summary", "")
            name = summary_response.get("name", "")
            email = summary_response.get("email", "")
            phone = summary_response.get("phone", "")

            candidate_id = self.vector_db.insert_row(
                table_name=self.table_name,
                summary=summary,
                full_resume=resume_text,
                email=email,
                phone=phone,
                name=name
            )

            # Save the file_bytes to the local folder
            with open(f"resumes/{candidate_id}{file_extension}", "wb") as f:
                f.write(file_bytes)
            return "Resume stored successfully"
        except Exception as e:
            return f"Error storing resume: {str(e)}"

    def retrieve_top_resumes(self, job_description: str, top_n: int = 10) -> List[tuple]:
        """Retrieve top matching resumes based on job description."""
        try:
            results = self.vector_db.retrieve_top_results(self.table_name, top_n, job_description)
            return results
        except Exception as e:
            print(f"Error retrieving resumes: {e}")
            return []

    def extract_criteria_from_jd(self, job_description: str) -> str:
        """Extract evaluation criteria from job description using LLM with Ericsson-specific context."""
        with open('prompts/extract_criteria.txt', 'r', encoding='utf-8') as f:
            prompt = f.read().format(job_description=job_description)

        return self.llm.get_chat_completion(prompt, "criteria_extractor")

    def summarize_resume(self, resume_text: str) -> str:
        """Generate a summary of the resume using LLM."""
        with open('prompts/summarize_resume.txt', 'r', encoding='utf-8') as f:
            prompt = f.read().replace("{resume_text}" , resume_text)
        response =  self.llm.get_chat_completion(prompt, "resume_analyzer")
        return response

    def evaluate_top_resumes(self, job_description: str, top_n: int = 10) -> List[Dict[str, Any]]:
        """Evaluate top resumes against job description and auto-extracted criteria."""
        top_resumes = self.retrieve_top_resumes(job_description, top_n)

        # Always extract criteria from job description
        criteria = self.extract_criteria_from_jd(job_description)

        # Prepare JSON input for LLM prompt
        resumes_json = []
        for result in top_resumes:
            resumes_json.append({
                "candidate_id": result.get("candidate_id"),
                "resume": result.get("metadata", {}).get("full_resume", "")
            })

        # Build the full JSON input for batch evaluation
        input_json = {
            "job_description": job_description,
            "resumes": resumes_json
        }

        # Load prompt template and inject data
        with open('prompts/evaluate_resume.txt', 'r', encoding='utf-8') as f:
            prompt_template = f.read()
        
        # Replace placeholders with actual data
        prompt = prompt_template.replace('{resume_text}', str(input_json))
        prompt = prompt.replace('{criteria}', str(criteria))
        
        # Send batch evaluation request to LLM
        evaluation_response = self.llm.get_chat_completion(prompt, "resume_evaluator")

        parsed_response = json.loads(evaluation_response)
        # Format results to match expected output
        formatted_results = []
        for i, eval_result in enumerate(parsed_response["results"]):
            # Get the candidate details from the database
            table_entry = self.vector_db.get_row_with_id(self.table_name,eval_result.get("candidate_id", ""))
            candidate_details = table_entry['metadatas'][0]
            formatted_results.append({
                "name": candidate_details.get("name", ""),
                "email": candidate_details.get("email", ""),
                "phone": candidate_details.get("phone", ""),
                "resume_id": i + 1,
                "candidate_id": eval_result.get("candidate_id", ""),
                "similarity_score": top_resumes[i]["similarity"] if i < len(top_resumes) else 0,
                "evaluation": eval_result.get("evaluation", {}),
                "resume_preview": top_resumes[i]["metadata"].get("full_resume", top_resumes[i]["summary"])[:200] + "..." if i < len(top_resumes) else ""
            })
        return formatted_results